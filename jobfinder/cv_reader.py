from pathlib import Path

from docx import Document


def read_cv(path: str | Path) -> str:
    """读取 .docx 或 .md 文件，返回纯文本。"""
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"CV 文件不存在：{path}")

    suffix = p.suffix.lower()
    if suffix == ".md":
        return p.read_text(encoding="utf-8")
    elif suffix == ".docx":
        return _read_docx(p)
    else:
        raise ValueError(f"不支持的文件格式：{suffix}，仅支持 .md 和 .docx")


def _read_docx(path: Path) -> str:
    doc = Document(str(path))
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # 提取表格内容
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)
